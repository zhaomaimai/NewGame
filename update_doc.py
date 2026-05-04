# version: C0_v1
# last_modified_cycle: C1
"""
Update 项目规划文档.docx with C0 + C1 completion status.
Creates a backup before modifying.
"""
import shutil
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC_PATH = "d:/BaiduNetdiskDownload/三国志5威力加强版/项目规划文档.docx"
BACKUP_PATH = "d:/BaiduNetdiskDownload/三国志5威力加强版/项目规划文档_backup_before_C1.docx"

def make_backup():
    if not os.path.exists(BACKUP_PATH):
        shutil.copy2(DOC_PATH, BACKUP_PATH)
        print(f"[DOC] Backup saved to: {BACKUP_PATH}")
    else:
        print(f"[DOC] Backup already exists: {BACKUP_PATH}")

def find_paragraph(doc, marker):
    """Find paragraph index containing marker text."""
    for i, p in enumerate(doc.paragraphs):
        if marker in p.text:
            return i
    return None

def update_document():
    doc = Document(DOC_PATH)

    # =============================================
    # 1. Update section 2: add backup rule
    # =============================================
    # Find "2.6 禁止行为" and add a new rule section after it
    idx_26 = find_paragraph(doc, "2.6 禁止行为")
    if idx_26 is not None:
        # Find the last paragraph of section 2.6 (next heading)
        insert_pos = idx_26 + 1
        while insert_pos < len(doc.paragraphs):
            text = doc.paragraphs[insert_pos].text.strip()
            if text.startswith("2.7") or text.startswith("3."):
                break
            insert_pos += 1

        # Add backup rule as a new paragraph at insert position
        p = doc.paragraphs[insert_pos].insert_paragraph_before(
            "2.7 文档更新规则（强制）"
        )
        p.style = doc.styles["Heading 2"]
        run = p.runs[0] if p.runs else p.add_run("2.7 文档更新规则（强制）")
        run.font.size = Pt(13)

        rules = [
            "【AI执行规范】 每个 Cycle 完成后必须更新本文档：",
            "更新前必须先复制一份本文档副本，文件命名为：项目规划文档_backup_before_C{当前Cycle编号}.docx",
            "然后修改本文档中以下内容：",
            "  - 第3章「当前系统状态」已完成模块列表",
            "  - 第3章「当前进度」状态表（⬜ → ✅）",
            "  - 更新对应 Cycle 的状态为 ✅ 已完成",
            "禁止：不更新文档直接进入下一 Cycle",
            "禁止：多人同时修改文档（单人顺序开发）",
        ]
        for rule in rules:
            rp = doc.paragraphs[insert_pos].insert_paragraph_before(rule)
            rp.style = doc.styles["List Bullet"]

        print("[DOC] Added section 2.7 backup rule")
    else:
        print("[DOC] Warning: section 2.6 not found")

    # =============================================
    # 2. Update section 3: completed modules
    # =============================================
    idx_completed = find_paragraph(doc, "已完成模块")
    if idx_completed is not None:
        # Find the table or list after "已完成模块"
        pos = idx_completed + 1
        # Skip existing status entries until we find the right spot
        while pos < len(doc.paragraphs):
            text = doc.paragraphs[pos].text.strip()
            if "当前GameState结构" in text or "当前进度" in text:
                break
            pos += 1

        # Insert C0 and C1 module entries before "当前GameState结构"
        if "当前GameState结构" in doc.paragraphs[pos].text:
            new_entries = [
                "C0-基础设施与调试系统 GameState/DebugSystem/RandomManager/TestRunner NewGame/scripts/core/ *.gd ✅ 稳定",
                "C1-数据管线 调色板/地图PNG/城市JSON导出 NewGame/export_*.py + data/ ✅ 稳定",
            ]
            for entry in reversed(new_entries):
                rp = doc.paragraphs[pos].insert_paragraph_before(entry)
                # Copy formatting from surrounding paragraphs
                rp.style = doc.styles["List Bullet"]
                for run in rp.runs:
                    run.font.size = Pt(10)
            print("[DOC] Updated completed modules")
    else:
        print("[DOC] Warning: completed modules section not found")

    # =============================================
    # 3. Update section 3: progress table
    # =============================================
    idx_progress = find_paragraph(doc, "当前进度")
    if idx_progress is not None:
        # Find the table rows - they should be in paragraphs after "当前进度"
        pos = idx_progress + 1
        row_count = 0
        while pos < len(doc.paragraphs) and row_count < 20:
            text = doc.paragraphs[pos].text.strip()
            # Look for C0 or C1 in the text
            if text.startswith("C0") or "C0基础设施" in text:
                # Replace ⬜ with ✅ and add "(C0_v1)"
                for run in doc.paragraphs[pos].runs:
                    if "⬜" in run.text:
                        run.text = run.text.replace("⬜", "✅")
                # Add status info
                doc.paragraphs[pos].add_run("  — GameState/Debug/Test/Random 模块已创建")
                print(f"[DOC] Updated C0 status at paragraph {pos}")
                row_count += 1
            elif text.startswith("C1") or "C1数据" in text:
                for run in doc.paragraphs[pos].runs:
                    if "⬜" in run.text:
                        run.text = run.text.replace("⬜", "✅")
                doc.paragraphs[pos].add_run("  — palette/map/city JSON导出完成")
                print(f"[DOC] Updated C1 status at paragraph {pos}")
                row_count += 1
            pos += 1
    else:
        print("[DOC] Warning: progress table not found")

    # =============================================
    # 4. Update Cycle C0 section
    # =============================================
    idx_c0 = find_paragraph(doc, "Cycle C0:")
    if idx_c0 is not None:
        pos = idx_c0 + 1
        while pos < len(doc.paragraphs):
            text = doc.paragraphs[pos].text.strip()
            if "状态：" in text:
                for run in doc.paragraphs[pos].runs:
                    if "⬜" in run.text:
                        run.text = run.text.replace("⬜", "✅") + " 已完成(C0_v1)"
                print(f"[DOC] Updated C0 status header")
                break
            pos += 1

        # Add completion summary after the cycle description
        pos = idx_c0 + 1
        while pos < len(doc.paragraphs):
            text = doc.paragraphs[pos].text.strip()
            if "下一Cycle提示词" in text:
                summary_lines = [
                    "",
                    "【已完成记录 C0_v1】",
                    "完成内容：",
                    "  - NewGame/project.godot: Godot 4.4+ 项目配置",
                    "  - NewGame/scripts/core/game_state.gd: GameState 单例 (autoload)",
                    "  - NewGame/scripts/core/debug_system.gd: DebugSystem (print_dbg + debug_mode)",
                    "  - NewGame/scripts/core/random_manager.gd: RandomManager (DEBUG_MODE 种子=42)",
                    "  - NewGame/scripts/core/test_runner.gd: TestRunner (register_test/run_all)",
                    "  - NewGame/scenes/debug/debug_panel.tscn: 调试面板场景",
                    "  - NewGame/scenes/debug/test_runner.tscn: 测试运行器场景",
                    "执行周期: 2026-05-04",
                ]
                for line in reversed(summary_lines):
                    rp = doc.paragraphs[pos].insert_paragraph_before(line)
                    rp.style = doc.styles["Normal"]
                print(f"[DOC] Added C0 completion summary")
                break
            pos += 1

    # =============================================
    # 5. Update Cycle C1 section
    # =============================================
    idx_c1 = find_paragraph(doc, "Cycle C1:")
    if idx_c1 is not None:
        pos = idx_c1 + 1
        while pos < len(doc.paragraphs):
            text = doc.paragraphs[pos].text.strip()
            if "状态：" in text:
                for run in doc.paragraphs[pos].runs:
                    if "⬜" in run.text:
                        run.text = run.text.replace("⬜", "✅") + " 已完成(C1_v1)"
                print(f"[DOC] Updated C1 status header")
                break
            pos += 1

        # Add completion summary
        pos = idx_c1 + 1
        while pos < len(doc.paragraphs):
            text = doc.paragraphs[pos].text.strip()
            if "下一Cycle提示词" in text:
                summary_lines = [
                    "",
                    "【已完成记录 C1_v1】",
                    "完成内容：",
                    "  - NewGame/export_palette.py: 调色板导出 (PAL256.S5 → data/palette.json, 256色)",
                    "  - NewGame/export_map_data.py: 地图PNG导出 (MAP256.S5 → data/map_full.png 800×592)",
                    "  - NewGame/export_city_data.py: 城市数据导出 (42 cities → data/cities.json)",
                    "  - 所有数据文件验证通过: JSON UTF-8编码, PNG正确渲染",
                    "执行周期: 2026-05-04",
                    "",
                    "【数据说明】",
                    "  - 调色板在PAL256.S5偏移768处, 直接读取RGB值(0-255), 无需×4缩放",
                    "  - 城市像素坐标使用硬编码映射表(已验证), SNDATA坐标与像素坐标不一致",
                    "  - MAP256.S5为800×592单字节索引图, PAL256.S5提供256色调色板",
                ]
                for line in reversed(summary_lines):
                    rp = doc.paragraphs[pos].insert_paragraph_before(line)
                    rp.style = doc.styles["Normal"]
                print(f"[DOC] Added C1 completion summary")
                break
            pos += 1

    # =============================================
    # 6. Update C2 下一Cycle提示词 with specific file paths
    # =============================================
    idx_c2 = find_paragraph(doc, "Cycle C2:")
    if idx_c2 is not None:
        pos = idx_c2 + 1
        while pos < len(doc.paragraphs):
            text = doc.paragraphs[pos].text.strip()
            if "下一Cycle提示词" in text:
                # Update the prompt to point to NewGame/ paths
                new_prompt = (
                    "【更新后的C2启动提示词】\n"
                    "前置条件:\n"
                    "- Cycle C0: NewGame/project.godot + GameState + DebugSystem + TestRunner 已创建\n"
                    "- Cycle C1: NewGame/data/map_full.png + NewGame/data/cities.json 已生成\n"
                    "请完成以下全部内容:\n"
                    "=== 第一步: 主场景 ===\n"
                    "创建 NewGame/scenes/main/main.tscn + main.gd\n"
                    "- 设置为 project.godot 的 main_scene\n"
                    "- _ready() 中加载 NewGame/data/cities.json → GameState.set_data('city.list', data)\n"
                    "- 实例化 MapView\n"
                    "=== 第二步: MapView ===\n"
                    "创建 NewGame/scenes/map/map_view.tscn + scripts/map/map_view.gd\n"
                    "- 使用 TextureRect 加载 map_full.png 作为地图\n"
                    "- 包裹在 SubViewport + Camera2D 中控制\n"
                    "- 鼠标拖拽平移 + 滚轮缩放(0.5x-3.0x)\n"
                    "- zoom 值存入 GameState.set_data('map.zoom', value)\n"
                    "=== 第三步: 城市标记 ===\n"
                    "创建 NewGame/scenes/map/city_marker.tscn + scripts/map/city_marker.gd\n"
                    "- 从 GameState.get_data('city.list') 获取城市列表\n"
                    "- 每个城市坐标处放置彩色圆点标记(直径8px)\n"
                    "- 鼠标悬停 → Tooltip 显示城市名\n"
                    "- 鼠标点击 → 发射 city_clicked(city_id) 信号\n"
                    "=== 版本控制 ===\n"
                    "每个新建文件头部标注: # version: C2_v1 / # last_modified_cycle: C2\n"
                    "注意: 所有路径相对 NewGame/ 目录"
                )
                rp = doc.paragraphs[pos].insert_paragraph_before(new_prompt)
                rp.style = doc.styles["Normal"]
                print(f"[DOC] Updated C2 startup prompt")
                break
            pos += 1

    # Save
    doc.save(DOC_PATH)
    print(f"\n[DOC] Document saved: {DOC_PATH}")
    print("[DOC] Update complete!")


if __name__ == "__main__":
    make_backup()
    update_document()
