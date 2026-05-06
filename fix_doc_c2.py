# version: C2_v2
# last_modified_cycle: C2
"""
Fix remaining issues in the planning document:
1. Update progress table C2 row: ⬜ → ✅
2. Clean up leftover "未开始" in status headers
"""
import os
import sys
from docx import Document

DOC_PATH = "d:/BaiduNetdiskDownload/三国志5威力加强版/项目规划文档.docx"


def fix_document():
    doc = Document(DOC_PATH)

    # =============================================
    # 1. Fix C2 status header: remove leftover "未开始"
    # =============================================
    for p in doc.paragraphs:
        if "状态：" in p.text:
            for run in p.runs:
                if "未开始" in run.text:
                    run.text = run.text.replace("未开始", "").strip()
                    print(f"[DOC] Cleaned '未开始' from: {run.text[:50]}")

    # =============================================
    # 2. Update progress TABLE (not paragraph)
    # =============================================
    for table in doc.tables:
        for row in table.rows:
            row_text = "|".join(cell.text for cell in row.cells)
            if "C2" in row_text:
                sys.stderr.write(f"[DOC] Found C2 in table: {row_text}\n")
                for cell in row.cells:
                    if "⬜" in cell.text or "未开始" in cell.text:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if "⬜" in run.text:
                                    run.text = run.text.replace("⬜", "✅")
                                if "未开始" in run.text:
                                    run.text = run.text.replace("未开始", " 已完成(C2_v1)")
                            # Also handle paragraph-level text
                            if "⬜" in p.text:
                                for run in p.runs:
                                    if "⬜" in run.text:
                                        run.text = run.text.replace("⬜", "✅")

    # Save
    doc.save(DOC_PATH)
    print(f"\n[DOC] Document saved: {DOC_PATH}")
    print("[DOC] Fix complete!")


if __name__ == "__main__":
    fix_document()
