# version: C2_v3
# last_modified_cycle: C2
"""
Update 项目规划文档.docx: mark C2 as in-progress, add accurate records.
Clean single-pass update.
"""
import shutil
import os
import sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

DOC_PATH = "d:/BaiduNetdiskDownload/三国志5威力加强版/项目规划文档.docx"
BACKUP_PATH = "d:/BaiduNetdiskDownload/三国志5威力加强版/项目规划文档_backup_before_C2v3.docx"


def make_backup():
    if not os.path.exists(BACKUP_PATH):
        shutil.copy2(DOC_PATH, BACKUP_PATH)
        sys.stderr.write(f"[DOC] Backup saved\n")


def replace_text(p, old, new):
    if old not in p.text:
        return False
    for run in p.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    return False


def clear_para(p):
    """Remove all runs and text from a paragraph."""
    for run in p.runs:
        run.text = ""
    # Also clear any remaining XML text nodes
    for t in p._element.findall(qn('w:r')):
        p._element.remove(t)


def set_para_text(p, text, size=Pt(11)):
    """Clear paragraph and set new text."""
    clear_para(p)
    run = p.add_run(text)
    run.font.size = size
    return run


def update_document():
    doc = Document(DOC_PATH)

    # ── 1. Update C2 section header status ──
    for i, p in enumerate(doc.paragraphs):
        if "Cycle C2:" in p.text:
            for j in range(i, min(i + 5, len(doc.paragraphs))):
                sp = doc.paragraphs[j]
                if "状态：" in sp.text:
                    replace_text(sp, "✅", "🔄")
                    replace_text(sp, "已完成(C2_v1)", "进行中(C2_v3)")
                    sys.stderr.write(f"[DOC] Status updated\n")
            break

    # ── 2. Update C2 goal description ──
    for i, p in enumerate(doc.paragraphs):
        if "Cycle C2:" in p.text:
            for j in range(i, min(i + 10, len(doc.paragraphs))):
                t = doc.paragraphs[j].text.strip()
                if "在 Godot 中显示完整" in doc.paragraphs[j].text:
                    replace_text(
                        doc.paragraphs[j],
                        "在 Godot 中显示完整 800×592 地图，支持鼠标拖拽平移、滚轮缩放。",
                        "在 Godot 中放置地图 + 城市标记 + 城市名称，支持城市位置编辑和连接编辑。"
                    )
                    sys.stderr.write(f"[DOC] Goal updated\n")
                    break
            break

    # ── 3. Update completed modules table (Table 1) ──
    table1 = doc.tables[1]
    for row in table1.rows:
        row_text = "".join(cell.text for cell in row.cells)
        if "C2-地图" in row_text:
            cell = row.cells[3]
            for p in cell.paragraphs:
                if "✅" in p.text:
                    replace_text(p, "✅", "🔄")
                    sys.stderr.write(f"[DOC] Modules table updated\n")
            break

    # ── 4. Update progress table (Table 2) ──
    table2 = doc.tables[2]
    for row in table2.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if cells[0] == "C2":
            cell = row.cells[2]
            for p in cell.paragraphs:
                replace_text(p, "✅", "🔄")
                replace_text(p, "已完成(C2_v1)", "进行中(C2_v3)")
                sys.stderr.write(f"[DOC] Progress table updated\n")
            break

    # ── 5. Replace the C2 completed record section ──
    # Find "【已完成记录 C2_v1】" and clear everything up to the separator
    # or "下一Cycle提示词" before it.
    # Also clear the old "【更新后的C2启动提示词】" block.
    record_start = -1
    for i, p in enumerate(doc.paragraphs):
        if "【已完成记录 C2" in p.text:
            record_start = i
            break

    if record_start >= 0:
        # Also find and clear the old "【更新后的C2启动提示词】" (before the record)
        for i in range(record_start - 1, max(record_start - 5, 0), -1):
            if "【更新后的C2启动提示词】" in doc.paragraphs[i].text:
                # Clear from this paragraph up to the separator before record_start
                clear_start = i
                # Find the separator before clear_start
                for k in range(clear_start - 1, max(clear_start - 3, 0), -1):
                    t = doc.paragraphs[k].text.strip()
                    if t.startswith("─") and len(t) > 10:
                        clear_start = k + 1  # Keep the separator
                        break

                # Clear paragraphs from clear_start to the separator after record_start
                clear_end = record_start
                for k in range(record_start, min(record_start + 15, len(doc.paragraphs))):
                    t = doc.paragraphs[k].text.strip()
                    if t.startswith("─") and len(t) > 10:
                        clear_end = k  # Keep the separator
                        break
                    if "Cycle C3:" in t:
                        clear_end = k
                        break

                sys.stderr.write(f"[DOC] Clearing paragraphs {clear_start} to {clear_end - 1}\n")
                for k in range(clear_start, clear_end):
                    clear_para(doc.paragraphs[k])
                record_start = clear_start  # Where we'll insert new content
                break

    # ── 6. Insert new C2 progress record ──
    if record_start >= 0:
        new_lines = [
            "【进行中记录 C2_v3】",
            "完成内容：",
            "  - NewGame/scenes/main/main.tscn + scripts/main/main.gd: 主场景，加载cities.json到GameState",
            "  - NewGame/scenes/map/map_view.tscn + scripts/map/map_view.gd: 简化版MapView，仅保留城市标记+CityEditor",
            "  - NewGame/scenes/map/city_marker.tscn + scripts/map/city_marker.gd: 城市标记，支持缩放同步",
            "  - NewGame/scripts/map/city_editor.gd: 城市编辑器(E键切换)，拖拽城市+编辑连接+保存到cities_custom.json",
            "  - NewGame/project.godot: 设置 main_scene，Godot 4.6.2 兼容性修复",
            "  - 用户重新生成 map_full.png(1463×1075) 并直接放入场景",
            "未完成(待C2继续)：",
            "  - 城市点上显示城市名称",
            "  - 正确布置所有城市的坐标(x, y)，确保与地图对齐",
            "  - 城市连接关系可视化",
            "当前已知问题：",
            "  - 城市坐标不准确，需要逐个在地图编辑器里调整",
            "  - 无城市名称标签，难以辨认",
            "执行周期: 2026-05-05",
        ]
        for idx, line in enumerate(new_lines):
            pi = record_start + idx
            if pi < len(doc.paragraphs):
                set_para_text(doc.paragraphs[pi], line)
            else:
                # Add new paragraph
                p = doc.add_paragraph(line)
                # Move it to the right position — tricky, so avoid this case
                sys.stderr.write(f"[DOC] Warning: ran out of paragraphs at line {idx}\n")

        # ── 7. Update the "下一Cycle提示词" section (should be after our new content) ──
        next_prompt_start = -1
        for k in range(record_start, min(record_start + 25, len(doc.paragraphs))):
            if "下一Cycle提示词" in doc.paragraphs[k].text:
                next_prompt_start = k
                break

        if next_prompt_start >= 0:
            # Clear the old prompt content (the "现在开始 Cycle C2..." block)
            clear_until = next_prompt_start + 1
            for k in range(next_prompt_start + 1, min(next_prompt_start + 10, len(doc.paragraphs))):
                t = doc.paragraphs[k].text.strip()
                if t.startswith("─") and len(t) > 10:
                    clear_until = k
                    break
                if "Cycle C3:" in t:
                    clear_until = k
                    break
                clear_para(doc.paragraphs[k])

            # Set new prompt content in the paragraph after "下一Cycle提示词"
            new_prompt = [
                "前置条件：",
                "- C0: Godot 项目 + GameState + 测试框架已创建",
                "- C1: game/data/map_full.png + game/data/cities.json 已生成",
                "- 地图由用户直接在场景中放置（MapView 已简化，无缩放/拖拽功能）",
                "",
                "请完成以下内容（延续 C2，不要进入 C3）：",
                "1. 在城市标记点上显示城市名称（Label 悬浮在点上方或内部）",
                "2. 正确布置所有 42 个城市的坐标：打开编辑器(E键)逐个拖拽城市到地图上正确位置",
                "3. 编辑城市之间的连接关系，确保道路网络正确",
                "4. 编辑器调整完成后，保存到 cities_custom.json",
                "5. 验证：打开游戏后所有城市名称可见，位置与地图对齐",
            ]
            insert_at = next_prompt_start + 1
            for idx, line in enumerate(new_prompt):
                pi = insert_at + idx
                if pi < len(doc.paragraphs):
                    set_para_text(doc.paragraphs[pi], line)
                else:
                    sys.stderr.write(f"[DOC] Warning: ran out of paragraphs at prompt line {idx}\n")

            sys.stderr.write(f"[DOC] Next-cycle prompt updated\n")
        else:
            sys.stderr.write(f"[DOC] WARNING: no '下一Cycle提示词' found after record\n")

    # Save
    doc.save(DOC_PATH)
    sys.stderr.write(f"\n[DOC] Done!\n")


if __name__ == "__main__":
    make_backup()
    update_document()
