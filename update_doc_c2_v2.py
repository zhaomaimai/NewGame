# version: C2_v2
# last_modified_cycle: C2
"""
Update 项目规划文档.docx with C2 completion status.
Handles both paragraphs and table cells properly.
"""
import shutil
import os
import sys
from docx import Document
from docx.shared import Pt

DOC_PATH = "d:/BaiduNetdiskDownload/三国志5威力加强版/项目规划文档.docx"
BACKUP_PATH = "d:/BaiduNetdiskDownload/三国志5威力加强版/项目规划文档_backup_before_C2.docx"


def make_backup():
    if not os.path.exists(BACKUP_PATH):
        shutil.copy2(DOC_PATH, BACKUP_PATH)
        sys.stderr.write(f"[DOC] Backup saved to: {BACKUP_PATH}\n")
    else:
        sys.stderr.write(f"[DOC] Backup already exists: {BACKUP_PATH}\n")


def replace_in_paragraph(p, old_text, new_text):
    """Replace text across multiple runs within a single paragraph."""
    full_text = p.text
    if old_text not in full_text:
        return False
    for run in p.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
    # If text is split across runs, rebuild
    if old_text in p.text:
        for run in p.runs:
            if run.text and run.text.strip():
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    return True
    return False


def update_document():
    doc = Document(DOC_PATH)

    # =============================================
    # 1. Update C2 section header status
    # =============================================
    for idx, p in enumerate(doc.paragraphs):
        if "Cycle C2:" in p.text or "Cycle C2：" in p.text:
            sys.stderr.write(f"[DOC] Found C2 section header\n")
            for i in range(idx, min(idx + 5, len(doc.paragraphs))):
                sp = doc.paragraphs[i]
                if "状态：" in sp.text or "状态:" in sp.text:
                    sys.stderr.write(f"[DOC] Found status line: {sp.text}\n")
                    if "⬜" in sp.text:
                        replace_in_paragraph(sp, "⬜", "✅")
                    if "未开始" in sp.text:
                        replace_in_paragraph(sp, "未开始", " 已完成(C2_v1)")
                    sys.stderr.write(f"[DOC] Updated status: {sp.text}\n")
                    break
            break

    # =============================================
    # 2. Update progress TABLE
    # =============================================
    for table in doc.tables:
        for row in table.rows:
            row_text = "".join(cell.text for cell in row.cells)
            if "C2" in row_text:
                sys.stderr.write(f"[DOC] Found C2 in progress table\n")
                for cell in row.cells:
                    for p in cell.paragraphs:
                        full_text = p.text
                        if "⬜" in full_text:
                            replace_in_paragraph(p, "⬜", "✅")
                            sys.stderr.write(f"[DOC] Table: replaced ⬜\n")
                        if "未开始" in full_text:
                            replace_in_paragraph(p, "未开始", " 已完成(C2_v1)")
                            sys.stderr.write(f"[DOC] Table: replaced 未开始\n")

    # =============================================
    # 3. Add completion summary under C2 section
    # =============================================
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("下一Cycle提示词"):
            # Check if this is immediately after the C2 section
            # Look backwards for C2 marker
            found_c2 = False
            for j in range(i - 1, max(i - 30, 0), -1):
                if "Cycle C2:" in doc.paragraphs[j].text or "C2启动" in doc.paragraphs[j].text or "更新后的C2" in doc.paragraphs[j].text:
                    found_c2 = True
                    break
                if "Cycle C3:" in doc.paragraphs[j].text or "Cycle C1:" in doc.paragraphs[j].text:
                    break
            if not found_c2:
                continue

            # Check if C2 completion already exists
            already_done = False
            for j in range(i - 1, max(i - 20, 0), -1):
                if "【已完成记录 C2" in doc.paragraphs[j].text:
                    already_done = True
                    sys.stderr.write(f"[DOC] C2 completion already exists, skipping\n")
                    break
            if already_done:
                break

            summary_lines = [
                "",
                "【已完成记录 C2_v1】",
                "完成内容：",
                "  - NewGame/scenes/main/main.tscn + scripts/main/main.gd: 主场景，加载cities.json到GameState",
                "  - NewGame/scenes/map/map_view.tscn + scripts/map/map_view.gd: 地图视图，鼠标右键拖拽平移+滚轮缩放(0.5x-3.0x)",
                "  - NewGame/scenes/map/city_marker.tscn + scripts/map/city_marker.gd: 城市标记，彩色圆点+悬停Tooltip+点击信号",
                "  - NewGame/project.godot: 设置 main_scene",
                "执行周期: 2026-05-05",
            ]
            for line in reversed(summary_lines):
                rp = doc.paragraphs[i].insert_paragraph_before(line)
                rp.style = doc.styles["Normal"]
            sys.stderr.write(f"[DOC] Added C2 completion summary\n")
            break

    # =============================================
    # 4. Update completed modules list
    # =============================================
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if "C1-数据管线" in text:
            new_entry = "C2-地图渲染 主场景/MapView/城市标记 NewGame/scenes/map/ + main/ ✅ 稳定"
            rp = doc.paragraphs[i].insert_paragraph_before(new_entry)
            rp.style = doc.styles["List Bullet"]
            sys.stderr.write(f"[DOC] Added C2 to completed modules\n")
            break

    # Save
    doc.save(DOC_PATH)
    sys.stderr.write(f"\n[DOC] Document saved: {DOC_PATH}\n")
    sys.stderr.write("[DOC] Update complete!\n")


if __name__ == "__main__":
    make_backup()
    update_document()
