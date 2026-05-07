#!/usr/bin/env python3
# version: C3_v1
# last_modified_cycle: C3
# Update 项目规划文档.docx with C3 completion status and integrate memory content.

import shutil, sys, os

sys.stdout = open(sys.stdout.fileno(), mode='w', encoding='utf-8', buffering=1)
sys.stderr = open(sys.stderr.fileno(), mode='w', encoding='utf-8', buffering=1)

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

DOCX = "项目规划文档.docx"
BACKUP = "项目规划文档_backup_before_C3.docx"

# ── helpers ──────────────────────────────────────────────────────

def find_para(doc, text, start=0):
    """Find index of paragraph containing `text`."""
    for i, p in enumerate(doc.paragraphs[start:], start=start):
        if text in p.text:
            return i
    return -1

def find_cell(doc, text):
    """Find (table_idx, row_idx, cell_idx) of a table cell containing `text`."""
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                if text in cell.text:
                    return ti, ri, ci
    return None

def set_para_text(para, text, bold=False, size=None, color=None):
    """Replace all runs in a paragraph with single-run text."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        run = para.runs[0]
        run.text = text
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
    else:
        run = para.add_run(text)
        run.bold = bold

def add_para_after(doc, idx, text, bold=False, size=None):
    """Insert a new paragraph after paragraph idx."""
    if idx + 1 < len(doc.paragraphs):
        new_p = doc.paragraphs[idx + 1].insert_paragraph_before(text)
    else:
        new_p = doc.add_paragraph(text)
    for run in new_p.runs:
        run.bold = bold
        if size:
            run.font.size = Pt(size)
    return new_p

# ── main ─────────────────────────────────────────────────────────

def main():
    # Backup
    if os.path.exists(BACKUP):
        print(f"Backup already exists: {BACKUP}")
    else:
        shutil.copy2(DOCX, BACKUP)
        print(f"Backup created: {BACKUP}")

    doc = Document(DOCX)
    modified = False

    # ── 1. Update C3 status in "已完成 Cycle" list (Section 3) ──
    idx = find_para(doc, "⬜ C3: 城市系统")
    if idx is not None:
        old = doc.paragraphs[idx].text
        new_text = "✅ C3: 城市系统 (CityManager, CityInfoPanel, connection_lines高亮, 可拖拽编辑面板, 防御力)"
        set_para_text(doc.paragraphs[idx], new_text)
        print(f"  [OK] Updated C3 status: {old} → {new_text}")
        modified = True
    else:
        print("  [!!] Could not find '⬜ C3: 城市系统' paragraph")

    # ── 2. Update progress table (C3 row: ⬜ 未开始 → ✅ 已完成) ──
    cell_loc = find_cell(doc, "⬜ 未开始")
    if cell_loc:
        ti, ri, ci = cell_loc
        cell = doc.tables[ti].rows[ri].cells[ci]
        for p in cell.paragraphs:
            if "⬜ 未开始" in p.text:
                set_para_text(p, "✅ 已完成")
                print(f"  [OK] Updated progress table C3: ⬜ 未开始 → ✅ 已完成")
                modified = True
                break

    # ── 3. Update C3 cycle section status line ──
    idx = find_para(doc, "状态：⬜ 未开始")
    if idx is not None:
        # Find the one right after "Cycle C3: 城市系统"
        c3_start = find_para(doc, "Cycle C3: 城市系统")
        if c3_start is not None and c3_start < idx < c3_start + 10:
            set_para_text(doc.paragraphs[idx], "状态：✅ 已完成(C3_v2)")
            print(f"  [OK] Updated C3 section status to ✅ 已完成(C3_v2)")
            modified = True

    # ── 4. Add C3 completion details after the C3 human description ──
    # Find the line after "【人类说明】" in C3 section, insert completion record
    idx = find_para(doc, "这是 UI 交互的第一步，内政/军事都从这里进入。")
    if idx is not None:
        # Check if already updated
        has_record = any("已完成记录" in p.text for p in doc.paragraphs[idx:idx+30])
        if not has_record:
            records = [
                "【已完成记录 C3_v2】",
                "完成内容：",
                "- NewGame/scripts/city/city_manager.gd: 城市选择管理器 (select_city/deselect_city/get_selected/get_city/get_all_cities)",
                "- NewGame/scripts/ui/city_info_panel.gd + .tscn: 可拖拽信息面板（坐标/人口/金/粮/士兵/防御，SpinBox实时编辑）",
                "- NewGame/scripts/city/test_city.gd: 测试函数 (test_city_selection, test_get_city)",
                "- NewGame/scripts/map/connection_lines.gd: 选中城市连接线橙色高亮（C3_v1升级）",
                "- NewGame/scripts/main/main.gd: 整合信号流, 非编辑模式点击城市→面板显示, 空白→取消, 编辑模式自动取消",
                "- NewGame/data/cities_custom.json: 66城市新增 defense=500 字段",
                "",
                "新增接口：",
                "- CityManager: select_city(id)/deselect_city()/get_selected()/get_city(id)/get_all_cities()",
                "- CityManager信号: city_selected(city_data)/city_deselected()",
                "- CityInfoPanel: show_city(data)/hide_panel(), 可拖拽, SpinBox编辑数值, 保存到文件",
                "",
                "信号流：",
                "非编辑模式点击城市 → marker.city_clicked → main._on_marker_clicked",
                "  → CityManager.select_city(id) → city_selected(city_data)",
                "    → CityInfoPanel.show_city() + marker.set_selected(true) + connection_lines高亮",
                "非编辑模式点击空白 → _unhandled_input → CityManager.deselect_city()",
                "    → CityInfoPanel.hide() + marker.set_selected(false) + 连接线取消高亮",
                "按下E进入编辑模式 → _process检测edit_mode变化 → 自动deselect_city()",
                "",
                "数据存储：",
                "- GameState city.selected: 当前选中城市ID",
                "- city.list中各城市的 population/gold/food/soldiers/defense 可通过面板编辑",
                "- 点'保存到文件'写入 cities_custom.json",
                "",
                "执行周期: 2026-05-07",
            ]
            # Insert after the "可替换性" paragraph or after current content
            insert_at = find_para(doc, "最小运行示例", idx)
            if insert_at is None:
                insert_at = idx + 5
            for line in reversed(records):
                if line == "":
                    continue
                add_para_after(doc, insert_at, line)
            print(f"  [OK] Added C3 completion records ({len(records)} lines)")
            modified = True
        else:
            print("  [--] C3 completion records already exist")
    else:
        print("  [!!] Could not find C3 human description anchor")

    # ── 5. Update 下一Cycle提示词 for C3 to point to C4 ──
    idx = find_para(doc, "现在开始 Cycle C3：城市系统。")
    if idx is not None:
        # Replace the entire prompt block with C4 pointer
        # Find the end marker (next ──── line or next Cycle)
        end_idx = find_para(doc, "────────────────────────────────────────────────────────────", idx)
        if end_idx is None or end_idx == idx:
            end_idx = find_para(doc, "Cycle C4:", idx)
        if end_idx is None:
            end_idx = idx + 30

        # Clear from idx to end_idx (replace with C4 prompt pointer)
        # We'll just replace the first line and remove the rest
        set_para_text(doc.paragraphs[idx],
            "下一Cycle提示词: 请从文档中复制 Cycle C4 的「下一Cycle提示词」作为新对话的第一条消息。",
            bold=False, size=10)

        # Clear intermediate paragraphs
        for ci in range(idx + 1, min(end_idx, len(doc.paragraphs))):
            p = doc.paragraphs[ci]
            if "现在开始 Cycle" in p.text:
                break
            if p.text.strip():
                try:
                    set_para_text(p, "")
                except:
                    pass
        print(f"  [OK] Updated C3→C4 next cycle prompt")
        modified = True

    # ── 6. Add CityManager to module interface table (Section 6) ──
    # Find the table with "GameState" / "MapView" rows
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 3 and "MapView" in cells[0].text and "city.list" in cells[2].text:
                # Found the interface table — add CityManager row
                has_cm = any("CityManager" in r.cells[0].text for r in table.rows)
                if not has_cm:
                    new_row = table.add_row()
                    new_row.cells[0].text = "CityManager"
                    new_row.cells[1].text = "select_city/deselect_city/get_selected/get_city/get_all"
                    new_row.cells[2].text = "city.selected, city.list"
                    print(f"  [OK] Added CityManager to interface table")
                    modified = True
                break

    # ── 7. Update GameState namespace description ──
    idx = find_para(doc, "GameState 将在 Cycle C0 中创建。当前无运行中的 Godot 项目。")
    if idx is not None:
        set_para_text(doc.paragraphs[idx],
            "GameState 已在 C0 中创建并持续使用。当前是运行中的 Godot 4.4+ 项目。",
            bold=False, size=10)
        print(f"  [OK] Updated GameState stale description")
        modified = True

    # ── 8. Add C0-C2 detailed file table after the "已完成模块" section ──
    idx = find_para(doc, "已完成模块")
    if idx is not None:
        # Check if detailed table already exists
        has_detail = any("game_state.gd" in p.text for p in doc.paragraphs[idx:idx+40])
        if not has_detail:
            details = [
                "",
                "【C0-C2 完成明细】（基于2026-05-07代码审计）",
                "",
                "C0 - 基础设施 (C0_v1):",
                "  game_state.gd — 单例autoload, get_data/set_data/register/clear/save_game/load_game",
                "  debug_system.gd — debug_mode静态布尔 + print_dbg静态函数",
                "  random_manager.gd — randf/randi_range/seed, DEBUG_MODE种子固定42",
                "  test_runner.gd — register_test/run_all/run_single",
                "  debug_panel.tscn — DEBUG_MODE复选框 + 日志输出",
                "  test_runner.tscn — 测试列表 + 运行按钮",
                "",
                "C1 - 数据管线 (C1_v1):",
                "  export_palette.py — PAL256.S5 → palette.json (256色RGB)",
                "  export_map_data.py — MAP256.S5 + PAL256.S5 → map_full.png (1280×940)",
                "  export_city_data.py — SNDATA.S5 → cities.json (42城, 硬编码坐标映射表)",
                "",
                "C2 - 地图渲染 (C2_v7 / C2_v10):",
                "  main.gd (C2_v7) — 加载cities_custom.json(优先) + drawn_lines.json → 实例化MapView",
                "  map_view.gd (C2_v7) — Control + Markers + CityEditor + ConnectionLines + FreehandLines",
                "  city_marker.gd (C2_v6) — City2.png (75×60) + 名称标签 + 发射city_clicked信号",
                "  city_editor.gd (C2_v10) — 7种模式(SELECT/ADD/CONNECT/DELETE/RENAME/DRAW/DELETE_LINE)",
                "  connection_lines.gd (C2_v9) — Catmull-Rom样条曲线 + 路径点 + 手绘抖动效果",
                "  freehand_lines.gd (C2_v7) — 手绘装饰线条渲染",
                "",
                "数据文件:",
                "  cities_custom.json — 66座城市, 坐标和连接已手动调整完毕",
                "  cities.json — 原始42城数据(备用)",
                "  drawn_lines.json — 保存的手绘线条",
                "",
                "C2编辑器快捷键: E开关编辑模式, Esc取消当前操作",
                "",
            ]
            insert_after = find_para(doc, "城市列表", idx)
            if insert_after is None:
                insert_after = idx + 10
            for line in reversed(details):
                add_para_after(doc, insert_after, line)
            print(f"  [OK] Added C0-C2 detail table ({len(details)} lines)")
            modified = True
        else:
            print("  [--] C0-C2 details already exist")

    # ── 9. Update version and date in header ──
    idx = find_para(doc, "v2.0")
    if idx is not None:
        set_para_text(doc.paragraphs[idx], "v3.0 ｜ 2026-05-07 ｜ Python 3 + Godot 4.4+")
        print(f"  [OK] Updated document version to v3.0")
        modified = True

    # ── Save ──
    if modified:
        doc.save(DOCX)
        print(f"\n✅ Document saved to {DOCX}")
    else:
        print(f"\n-- No changes made.")

if __name__ == "__main__":
    main()
