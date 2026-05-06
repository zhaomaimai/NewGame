# version: C2_v1
# last_modified_cycle: C2
"""
Update 项目规划文档.docx with C2 completion status.
Creates a backup before modifying.
"""
import shutil
import os
from docx import Document

DOC_PATH = "d:/BaiduNetdiskDownload/三国志5威力加强版/项目规划文档.docx"
BACKUP_PATH = "d:/BaiduNetdiskDownload/三国志5威力加强版/项目规划文档_backup_before_C2.docx"


def make_backup():
    if not os.path.exists(BACKUP_PATH):
        shutil.copy2(DOC_PATH, BACKUP_PATH)
        print(f"[DOC] Backup saved to: {BACKUP_PATH}")
    else:
        print(f"[DOC] Backup already exists: {BACKUP_PATH}")


def find_paragraph(doc, marker):
    for i, p in enumerate(doc.paragraphs):
        if marker in p.text:
            return i
    return None


def update_document():
    doc = Document(DOC_PATH)

    # =============================================
    # 1. Update C2 status header and progress table
    # =============================================
    idx_c2 = find_paragraph(doc, "Cycle C2:")
    if idx_c2 is not None:
        pos = idx_c2 + 1
        while pos < len(doc.paragraphs):
            text = doc.paragraphs[pos].text.strip()
            if "状态：" in text:
                for run in doc.paragraphs[pos].runs:
                    if "⬜" in run.text:
                        run.text = run.text.replace("⬜", "✅") + " 已完成(C2_v1)"
                print(f"[DOC] Updated C2 status header")
                break
            pos += 1

        pos = idx_c2 + 1
        while pos < len(doc.paragraphs):
            text = doc.paragraphs[pos].text.strip()
            if "下一Cycle提示词" in text:
                # Check if already has completion record
                has_record = False
                check_pos = pos - 1
                while check_pos > 0 and doc.paragraphs[check_pos].text.strip():
                    if "【已完成记录 C2_v1】" in doc.paragraphs[check_pos].text:
                        has_record = True
                        break
                    check_pos -= 1

                if not has_record:
                    summary_lines = [
                        "",
                        "【已完成记录 C2_v1】",
                        "完成内容：",
                        "  - NewGame/scenes/main/main.tscn + scripts/main/main.gd: 主场景，加载cities.json到GameState",
                        "  - NewGame/scenes/map/map_view.tscn + scripts/map/map_view.gd: 地图视图，鼠标右键拖拽平移+滚轮缩放(0.5x-3.0x)",
                        "  - NewGame/scenes/map/city_marker.tscn + scripts/map/city_marker.gd: 城市标记，彩色圆点+悬停Tooltip+点击信号",
                        "  - NewGame/project.godot: 设置 main_scene",
                        "执行周期: 2026-05-05",
                    ]
                    for line in reversed(summary_lines):
                        rp = doc.paragraphs[pos].insert_paragraph_before(line)
                        rp.style = doc.styles["Normal"]
                    print(f"[DOC] Added C2 completion summary")
                break
            pos += 1

    # =============================================
    # 2. Update progress table: C2 ⬜ → ✅
    # =============================================
    idx_progress = find_paragraph(doc, "当前进度")
    if idx_progress is not None:
        pos = idx_progress + 1
        while pos < len(doc.paragraphs):
            text = doc.paragraphs[pos].text.strip()
            if "C2" in text and "地图渲染" in text:
                for run in doc.paragraphs[pos].runs:
                    if "⬜" in run.text:
                        run.text = run.text.replace("⬜", "✅")
                doc.paragraphs[pos].add_run("  — main.tscn + MapView拖拽/缩放 + 城市标记(悬停/点击)")
                print(f"[DOC] Updated C2 progress table")
                break
            pos += 1

    # =============================================
    # 3. Update C3 startup prompt with C2 context
    # =============================================
    idx_c3 = find_paragraph(doc, "Cycle C3:")
    if idx_c3 is not None:
        pos = idx_c3 + 1
        while pos < len(doc.paragraphs):
            text = doc.paragraphs[pos].text.strip()
            if "下一Cycle提示词" in text:
                new_prompt = (
                    "【更新后的C3启动提示词】\n"
                    "前置条件:\n"
                    "- Cycle C0: NewGame/project.godot + GameState + DebugSystem + TestRunner 已创建\n"
                    "- Cycle C1: NewGame/data/map_full.png + NewGame/data/cities.json 已生成\n"
                    "- Cycle C2: 地图可显示，鼠标拖拽/滚轮缩放，城市标记点击发射 city_clicked 信号\n"
                    "请完成以下全部内容:\n"
                    "=== 第一步: CityManager ===\n"
                    "创建 NewGame/scripts/city/city_manager.gd（普通 Node，非 autoload）\n"
                    "- func select_city(id: int) → void\n"
                    "- func get_selected() → Dictionary（返回城市数据，或 null）\n"
                    "- func get_city(id: int) → Dictionary\n"
                    "- func get_all_cities() → Array\n"
                    "- 数据存储在 GameState：set_data('city.selected', id) / get_data('city.list')\n"
                    "- 发射信号 city_selected(city_data) / city_deselected()\n"
                    "=== 第二步: 城市信息面板 ===\n"
                    "创建 NewGame/scenes/ui/city_info_panel.tscn + scripts/ui/city_info_panel.gd\n"
                    "- 在屏幕右侧显示，宽度约 300px\n"
                    "- 显示：城市名、坐标、人口、金、粮、士兵（用占位值）\n"
                    "- 显示连接城市名称列表\n"
                    "- 选中城市时显示，取消选中时隐藏\n"
                    "- 数据来源：GameState.get_data('city.list')[id]\n"
                    "=== 第三步: 连接线 ===\n"
                    "创建 NewGame/scenes/map/connection_lines.gd（附着在 MapView 上的脚本）\n"
                    "- 选中城市后，从该城市到所有连接城市绘制半透明直线\n"
                    "- 使用 _draw() + draw_line() 或 Line2D 节点\n"
                    "- 颜色：白色半透明，宽度 2px\n"
                    "=== 第四步: 连接 city_marker 信号 ===\n"
                    "在 main.gd 中连接 city_marker 的 city_clicked → CityManager.select_city()\n"
                    "=== 版本控制 ===\n"
                    "每个新建文件头部标注: # version: C3_v1 / # last_modified_cycle: C3\n"
                    "注意: 所有路径相对 NewGame/ 目录"
                )
                rp = doc.paragraphs[pos].insert_paragraph_before(new_prompt)
                rp.style = doc.styles["Normal"]
                print(f"[DOC] Updated C3 startup prompt")
                break
            pos += 1

    # =============================================
    # 4. Update C2 module entry in completed modules
    # =============================================
    idx_completed = find_paragraph(doc, "已完成模块")
    if idx_completed is not None:
        pos = idx_completed + 1
        while pos < len(doc.paragraphs):
            text = doc.paragraphs[pos].text.strip()
            if "当前GameState结构" in text or "C1-数据管线" in text:
                break
            pos += 1

        if "C1-数据管线" in doc.paragraphs[pos].text:
            new_entries = [
                "C2-地图渲染 主场景/MapView/城市标记 NewGame/scenes/map/ + main/ ✅ 稳定",
            ]
            for entry in reversed(new_entries):
                rp = doc.paragraphs[pos].insert_paragraph_before(entry)
                rp.style = doc.styles["List Bullet"]
            print("[DOC] Updated completed modules")
        elif "当前GameState结构" in doc.paragraphs[pos].text:
            # Insert before "当前GameState结构"
            new_entries = [
                "C2-地图渲染 主场景/MapView/城市标记 NewGame/scenes/map/ + main/ ✅ 稳定",
            ]
            for entry in reversed(new_entries):
                rp = doc.paragraphs[pos].insert_paragraph_before(entry)
                rp.style = doc.styles["List Bullet"]
            print("[DOC] Updated completed modules")

    # Save
    doc.save(DOC_PATH)
    print(f"\n[DOC] Document saved: {DOC_PATH}")
    print("[DOC] Update complete!")


if __name__ == "__main__":
    make_backup()
    update_document()
