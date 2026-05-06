# version: C2_v6
# last_modified_cycle: C2
"""
Update 项目规划文档.docx: update C2 progress record to C2_v6.
"""
import shutil
import os
import sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

DOC_PATH = "d:/BaiduNetdiskDownload/三国志5威力加强版/项目规划文档.docx"
BACKUP_PATH = "d:/BaiduNetdiskDownload/三国志5威力加强版/项目规划文档_backup_before_C2v6.docx"


def make_backup():
    if not os.path.exists(BACKUP_PATH):
        shutil.copy2(DOC_PATH, BACKUP_PATH)
        sys.stderr.write(f"[DOC] Backup saved\n")


def clear_para(p):
    for run in p.runs:
        run.text = ""
    for t in p._element.findall(qn('w:r')):
        p._element.remove(t)


def set_para_text(p, text, size=Pt(11)):
    clear_para(p)
    if text:
        run = p.add_run(text)
        run.font.size = size
    return p


def find_para_containing(doc, keyword, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if keyword in p.text:
            return i
    return -1


def update_document():
    doc = Document(DOC_PATH)

    # ── 1. Update C2 status ──
    for i, p in enumerate(doc.paragraphs):
        if "Cycle C2:" in p.text:
            for j in range(i, min(i + 5, len(doc.paragraphs))):
                sp = doc.paragraphs[j]
                if "C2_v3" in sp.text:
                    sp.text = sp.text.replace("C2_v3", "C2_v6")
                    sys.stderr.write(f"[DOC] Status updated\n")
            break

    # ── 2. Update goal description ──
    for i, p in enumerate(doc.paragraphs):
        if "Cycle C2:" in p.text:
            for j in range(i, min(i + 10, len(doc.paragraphs))):
                t = doc.paragraphs[j].text.strip()
                if "在 Godot 中显示" in t or "在 Godot 中放置" in t:
                    old_text = "在 Godot 中放置地图 + 城市标记 + 城市名称，支持城市位置编辑和连接编辑。"
                    new_text = "1280×940 地图 + 城池贴图(City2.png) + 城市名称标签，支持可视化编辑器(添加/移动/连接/删除/改名)。"
                    doc.paragraphs[j].text = doc.paragraphs[j].text.replace(old_text, new_text)
                    sys.stderr.write(f"[DOC] Goal updated\n")
                    break
            break

    # ── 3. Find the record_start and next_prompt_start BEFORE clearing ──
    record_start = find_para_containing(doc, "【进行中记录 C2")
    if record_start < 0:
        sys.stderr.write(f"[DOC] ERROR: could not find C2 progress record\n")
        return

    # Find "下一Cycle提示词" AFTER the current record
    orig_next_start = find_para_containing(doc, "下一Cycle提示词", start=record_start + 5)
    if orig_next_start < 0:
        sys.stderr.write(f"[DOC] ERROR: could not find 下一Cycle提示词 after record\n")
        return

    # Find the clear end: either separator, Cycle C3 header, or 下一Cycle提示词
    clear_end = record_start
    for k in range(record_start, min(record_start + 30, len(doc.paragraphs))):
        t = doc.paragraphs[k].text.strip()
        if t.startswith("─") and len(t) > 10:
            clear_end = k  # Keep separator
            break
        if "下一Cycle提示词" in t:
            clear_end = k - 1  # Stop BEFORE 下一Cycle提示词 header
            break
        if "Cycle C3:" in t:
            clear_end = k
            break
        clear_end = k + 1  # Include this paragraph in clear range

    sys.stderr.write(f"[DOC] record_start={record_start}, orig_next_start={orig_next_start}, clear_end={clear_end}\n")

    # Clear old record content (not including 下一Cycle提示词 header)
    for k in range(record_start, min(clear_end, len(doc.paragraphs))):
        clear_para(doc.paragraphs[k])

    # ── 4. Insert new C2 progress record ──
    new_lines = [
        "【进行中记录 C2_v6】",
        "完成内容：",
        "  - 地图显示: 用户将 map_full.png(1280×940) 直接在 main.tscn 中放置为 Sprite2D",
        "  - 城池显示: city_marker.tscn + city_marker.gd 重写，显示 City2.png 贴图(75×60) + 城市名称标签",
        "  - 坐标系统: 从 800×592 参考空间改为 1280×940 直接像素坐标，42城坐标已换算",
        "  - 编辑器重写: city_editor.gd，5种模式(选择/添加/连接/删除/改名)+拖拽+改名的弹出面板",
        "  - 工具栏: 模式按钮 + 保存/加载 + 2秒自动消失的反馈提示",
        "  - 连接线: 编辑模式下显示连接线，选中城市黄色高亮；连接模式绿色预览线",
        "  - map_view.tscn 简化: 去掉缩放/拖拽系统，去掉 REF_SIZE 换算,注释掉地图TextureRect",
        "未完成(待C2继续)：",
        "  - 逐个手动拖拽42个城市到地图上的正确位置",
        "  - 编辑城市之间的连接关系",
        "  - 保存最终配置到 cities_custom.json",
        "执行周期: 2026-05-06",
    ]

    # Calculate how many lines we can write without hitting orig_next_start
    max_lines = orig_next_start - record_start
    sys.stderr.write(f"[DOC] Writing {len(new_lines)} lines into {max_lines} available slots\n")

    for idx, line in enumerate(new_lines):
        pi = record_start + idx
        if pi < orig_next_start and pi < len(doc.paragraphs):
            set_para_text(doc.paragraphs[pi], line)
        else:
            sys.stderr.write(f"[DOC] Warning: ran out of space at line {idx} (pi={pi}, limit={orig_next_start})\n")

    # ── 5. Update the 下一Cycle提示词 content ──
    # orig_next_start still points to the "下一Cycle提示词" header paragraph
    # Clear old content after it
    clear_until = orig_next_start + 1
    for k in range(orig_next_start + 1, min(orig_next_start + 15, len(doc.paragraphs))):
        t = doc.paragraphs[k].text.strip()
        if t.startswith("─") and len(t) > 10:
            clear_until = k
            break
        if "Cycle C3:" in t:
            clear_until = k
            break
        clear_para(doc.paragraphs[k])

    # Set new prompt content
    new_prompt = [
        "前置条件：",
        "- C0: Godot 项目 + GameState + 测试框架已创建",
        "- C1: data/map_full.png + data/cities.json 已生成",
        "- C2: 地图 + 城池贴图(City2.png) + 编辑器已就绪",
        "",
        "请完成以下内容（延续 C2，不要进入 C3）：",
        "1. 逐个调整42个城市坐标：按E进入编辑模式 → 选择模式 → 拖拽城市到地图正确位置",
        "2. 编辑城市之间的连接关系：连接模式 → 点击两个城市切换连接",
        "3. 使用改名模式给缺少名称或默认名'新城'的城市重新命名",
        "4. 需要更多城市时使用添加模式创建",
        "5. 调整完成后点击保存，数据写入 cities_custom.json",
        "6. 验证：打开游戏后所有城市贴图位置与地图对齐，名称正确",
    ]
    insert_at = orig_next_start + 1
    for idx, line in enumerate(new_prompt):
        pi = insert_at + idx
        if pi < len(doc.paragraphs):
            set_para_text(doc.paragraphs[pi], line)
        else:
            sys.stderr.write(f"[DOC] Warning: ran out of paragraphs at prompt line {idx}\n")

    sys.stderr.write(f"[DOC] Next-cycle prompt updated\n")

    # Save
    doc.save(DOC_PATH)
    sys.stderr.write(f"\n[DOC] Done!\n")


if __name__ == "__main__":
    make_backup()
    update_document()
